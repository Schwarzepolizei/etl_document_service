from io import BytesIO
from docx import Document


def parse_docx(file_bytes: bytes) -> dict:
    """
    Возвращает структуру документа:
    {
        "full_text": str,
        "elements": [
            {"type": "paragraph", "text": "..."},
            {"type": "table_row", "text": "...", "cells": ["...", "..."]},
            ...
        ]
    }
    """
    doc = Document(BytesIO(file_bytes))

    elements = []
    full_text_parts = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            elements.append(
                {
                    "type": "paragraph",
                    "text": text,
                }
            )
            full_text_parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                cells.append(cell_text)

            if any(cell for cell in cells):
                row_text = " | ".join(cell if cell else "" for cell in cells).strip()

                elements.append(
                    {
                        "type": "table_row",
                        "text": row_text,
                        "cells": cells,
                    }
                )
                full_text_parts.append(row_text)

    return {
        "full_text": "\n\n".join(full_text_parts).strip(),
        "elements": elements,
    }